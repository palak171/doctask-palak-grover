"""
Format detection and text extraction. "It understands the pile" starts here:
mixed formats go in, plain text with clause boundaries comes out.

We deliberately keep this dependency-light (pypdf, python-docx) rather than
pulling in a heavy document-AI SDK — the brief asks for a system that a
stranger can run in minutes, and OCR-grade parsing is out of scope for a task
this size. That trade-off is logged here and in PROGRESS.md.
"""
from __future__ import annotations

import hashlib
import io
from dataclasses import dataclass


@dataclass
class ParsedDocument:
    text: str
    format: str


def sha256_of(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def parse(filename: str, data: bytes) -> ParsedDocument:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return ParsedDocument(text=_parse_pdf(data), format="pdf")
    if lower.endswith(".docx"):
        return ParsedDocument(text=_parse_docx(data), format="docx")
    if lower.endswith(".md"):
        return ParsedDocument(text=data.decode("utf-8", errors="replace"), format="md")
    # default: treat as plain text
    return ParsedDocument(text=data.decode("utf-8", errors="replace"), format="txt")


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _parse_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in document.paragraphs)
